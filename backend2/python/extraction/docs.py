from docx import Document

def extract_text_and_tables_from_docx(docx_path):
    document = Document(docx_path)
    all_text = ""
    all_tables = []
    for para in document.paragraphs:
        all_text += para.text + "\n"
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        all_tables.append(table_data)
    return all_text, all_tables
